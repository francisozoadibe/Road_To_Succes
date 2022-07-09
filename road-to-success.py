from docx import Document
from docx.shared import Inches
import pyttsx3


"""

in other to make this project available for other people to see
it and download it or experiment with it we have to done thing.

    1. uninstall the packages or libraries you install for this 
    projects.
        i. pip3 uninstall python-docx
            a. y for yes
        ii. pip3 uninstall pyttsx3
            b. yes
        iii. when you run the program it will tell you module not found
        and the reason why you have to remove the two packages by 
        uninstalling it manually is make it easier for the user of
        the same program not to download those packages manually. 
        But maket eaier for them to install the package with one
        command.
        To do this we have to create a file called "requirements.txt"
        with dot txt file extension
        
    2. open up your exploerer taB, create a new file name
        requirements with .txt file extension
        i. type the packages with their version inside the 
        requirement
            a. python-docx==0.8.11
            b. pyttsx3==2.90
    close the requirement file
    
    3. Go your terminal
        a. type "pip3 install -r requirements.txt and press enter
        you will see that python is using the dependencies inside
        requirement to download the required packages
        
        " we can see that, the package install is using the 
        requirements dependencies to install the dependencies
"""

def speak(text):
    pyttsx3.speak(text)


# Text to speech
"""
to use the text to speech library, you need instal the text to
speech library using this steps:
    1. go to your terminal and use the package manager pip3
    2. type "pip3 install pyttsx3"
    3. go up and import the pyttsx3 library
    
"""

# ADDING of More work experiences to your CV
"""
        our great works will be within or between creation of document 
        and saving of document       
    1.  document = Document()   
        and saving of document       
    2. document.save('cv.docx')
    
"""

document = Document() # this is how we create document

""" 
 what we will do here is to add work experience
"""

#PROFILE PICTURE
document.add_picture('ProPic.jpg', width=Inches(0.5))


# NAME, PHONE NUMBERS AND EMILS
name = input('what is your name?')
speak('hello' + name + 'how are you today?')

speak('what is ypur phone number?')
phone_number = input('what is your phone number?')
email = input('what is your email address?')

# ABOUT ME
document.add_heading('About Me')
document.add_paragraph(input('Tell me about your self'))

# Work Experience
document.add_heading(' Work Experience')

# to add more experiences we will use while loop
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date')
        to_date = input('To Date')
        # we are going to use p paragraph above to another paragraph below
        p.add_run(company + ' ').bold = True  # this how you things inside paragraph
        p.add_run(from_date + '_' + to_date + '\n').italic  = True

        # we can capture the work experience description below
        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
        # statement will keep on asking us quetion forever i.e infinity
    else:
        break    
    
# Skill
document.add_heading(' skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# to add more experiences we will use while loop
while True:
    has_more_skills = input('Do you have more SKILLS? Yes or No ')
    if has_more_skills.lower() == 'yes':      
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break   
# foter 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Chis and knowledge books'


document.save('cv.docx') # this how we save new document name
